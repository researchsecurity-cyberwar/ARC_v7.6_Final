# -*- coding: utf-8 -*-
"""
ArcChatEngine — mesin dialog ARC yang menyatukan modul-modul lain.

Jembatan antara manusia dan ARC untuk tukar informasi bug bounty:

    ConversationEngine      -> dialog stateful + memori per target
    CommandInterpreter      -> bahasa alami -> aksi operasional
    program_brief           -> parse/simpan/merge brief per target + manifest
    GoogleVRPIntegrator     -> intel program Google VRP (lazy, opsional)

Prinsip ketahanan (aman di Kali Linux):
- Semua import berat dilakukan LAZY (di dalam method) dengan try/except.
- Semua handler dibungkus try/except -> tidak pernah crash, selalu balas pesan.
- Tanpa dependency tambahan (stdlib only) selain modul ARC yang sudah ada.
"""
import json
import os
import re
import threading
from datetime import datetime
from typing import Any, Dict, List, Optional

from DIALOGIC_COPILLOT.conversation_engine import ConversationEngine
from DIALOGIC_COPILLOT.command_interpreter import CommandInterpreter
from DIALOGIC_COPILLOT.program_brief import (
    parse_program_brief,
    ProgramBriefStore,
    brief_to_manifest,
    _clean,
    _split_items,
)

MEMORY_DIR_DEFAULT = "~/.arc/conversations"
BRIEFS_DIR_DEFAULT = "~/.arc/briefs"
DEFAULT_TARGET = "umum"


class ArcChatEngine:
    """Facade dialog: brief bounty + intel + bantuan untuk semua target."""

    def __init__(self, memory_dir: Optional[str] = None,
                 briefs_dir: Optional[str] = None,
                 default_target: str = DEFAULT_TARGET,
                 enable_google_vrp: bool = True):
        self.memory_dir = os.path.expanduser(memory_dir or MEMORY_DIR_DEFAULT)
        self.briefs_dir = os.path.expanduser(briefs_dir or BRIEFS_DIR_DEFAULT)
        self.default_target = default_target
        self.enable_google_vrp = enable_google_vrp

        self.conversation = ConversationEngine(memory_dir=self.memory_dir)
        self.commands = CommandInterpreter()
        self.brief_store = ProgramBriefStore(briefs_dir=self.briefs_dir)
        self._google_integrator: Any = None
        self._google_error: Optional[str] = None
        self._lock = threading.RLock()

        # Wire: ConversationEngine -> brief handler (modul program_brief)
        self.conversation.set_brief_engine(self)

    # ==================================================================
    # PUBLIC API
    # ==================================================================
    def handle(self, message: str) -> str:
        """Proses satu pesan dari manusia -> jawaban ARC (thread-safe)."""
        with self._lock:
            message = (message or "").strip()
            if not message:
                return "Halo! Ketik /bantuan untuk daftar perintah."

            # 1) Perintah siklus hidup percakapan (mulai/ganti target)
            started = self._maybe_start_conversation(message)
            if started is not None:
                self._record(message, started)
                return started

            # 2) Pastikan ada percakapan aktif
            if self.conversation.current_conversation is None:
                target = self._extract_target(message) or self.default_target
                self.conversation.start_conversation(target_domain=target)

            # 3) Teruskan ke mesin dialog (brief_engine hook memproses command brief)
            conv = self.conversation.current_conversation
            before = len(conv["messages"]) if conv else 0
            try:
                self.conversation.add_message("human", message)
            except Exception as e:
                return f"⚠️ Terjadi kesalahan memproses pesan (safe): {e}"

            # 4) Baca ulang percakapan AKTIF (handler boleh mengganti percakapan)
            current = self.conversation.current_conversation
            if current is not None:
                new_msgs = current.get("messages", [])
                if current is conv:
                    new_msgs = new_msgs[before:]
                for msg in reversed(new_msgs):
                    if msg.get("role") == "arc" and msg.get("content"):
                        return msg["content"]
            return self._fallback_reply(message)

    def start_conversation(self, target: str, program_name: Optional[str] = None) -> str:
        """Mulai dialog untuk target apa saja."""
        with self._lock:
            target = (target or "").strip().lower() or self.default_target
            self.conversation.start_conversation(target_domain=target)
            msg = f"🎯 Dialog dimulai untuk target: {target}"
            if program_name:
                msg += f"\n📋 Program: {program_name}"
            msg += ("\nKirim detail bounty (deskripsi / scope / rules / requirements) "
                    "atau ketik /bantuan.")
            self._record(f"/mulai {target}", msg)
            return msg

    def get_history(self, limit: int = 200) -> List[Dict[str, Any]]:
        conv = self.conversation.current_conversation
        if conv is None:
            return []
        return conv.get("messages", [])[-limit:]

    def list_briefs(self) -> List[str]:
        return self.brief_store.list_briefs()

    # ==================================================================
    # SIKLUS HIDUP PERCAKAPAN
    # ==================================================================
    def _maybe_start_conversation(self, message: str) -> Optional[str]:
        low = message.lower().strip()
        m = re.match(r'^(?:/mulai|mulai|/start|start|target baru)\s+(?:target\s+)?(\S+)$', low)
        if m:
            target = m.group(1).strip().lower()
            self.conversation.start_conversation(target_domain=target)
            return (f"🎯 Dialog dimulai untuk target: {target}\n"
                    f"Kirim detail bounty atau /bantuan.")
        m = re.match(r'^(?:ganti target|switch target|pindah target)\s+(\S+)$', low)
        if m:
            target = m.group(1).strip().lower()
            self.conversation.start_conversation(target_domain=target)
            return (f"🔄 Target diganti: {target}\n"
                    f"Brief lama tetap tersimpan (lihat /daftar brief).")
        # "target X" singkat -> pindah target (ditangkap SEBELUM add_message
        # agar start_conversation tidak mereset percakapan di tengah proses)
        m = re.match(r'^target\s+(\S+)$', low)
        if m:
            target = m.group(1).lower()
            self.conversation.start_conversation(target_domain=target)
            return f"🎯 Target diganti: {target}\nKirim detail bounty atau /bantuan."
        return None

    def _extract_target(self, text: str) -> Optional[str]:
        m = re.search(r'(?:target|target_domain|domain|host)\s*[:=]\s*([a-zA-Z0-9._\-/]+)',
                      text, re.IGNORECASE)
        if m:
            return m.group(1).strip().lower()
        # "buat brief untuk X, scope ..." / "for X"
        m = re.search(r'\b(?:untuk|for)\s+([a-zA-Z0-9._\-/]+)', text, re.IGNORECASE)
        if m:
            return m.group(1).strip().lower()
        return None

    def _record(self, human_msg: str, arc_resp: str):
        """Catat percakapan langsung (tanpa auto-process ulang)."""
        conv = self.conversation.current_conversation
        if conv is None:
            return
        now = datetime.now().isoformat()
        conv["messages"].append({"role": "human", "content": human_msg,
                                 "timestamp": now, "metadata": {}})
        conv["messages"].append({"role": "arc", "content": arc_resp,
                                 "timestamp": now, "metadata": {}})
        try:
            self.conversation._save_conversation()
        except Exception:
            pass

    def _last_arc_reply(self) -> Optional[str]:
        conv = self.conversation.current_conversation
        if conv is None:
            return None
        for msg in reversed(conv.get("messages", [])):
            if msg.get("role") == "arc" and msg.get("content"):
                return msg["content"]
        return None


    # ==================================================================
    # HANDLER PERINTAH BRIEF (dipanggil ConversationEngine via hook)
    # ==================================================================
    def handle_brief_command(self, message: str) -> Optional[str]:
        """Proses perintah terkait brief bounty. None = bukan perintah brief."""
        try:
            message = (message or "").strip()
            if not message:
                return None
            low = message.lower()
            conv = self.conversation.current_conversation
            target = conv["target_domain"] if conv else None

            # --- bantuan / help ---
            if low in ("/bantuan", "/help", "bantuan", "help", "?", "tolong"):
                return self.help_text

            # --- info target ---
            if low in ("target", "/target", "apa target"):
                return f"🎯 Target aktif: {target or 'belum ada'}"

            # --- daftar brief ---
            if low in ("daftar brief", "list briefs", "/briefs", "semua brief"):
                return self._fmt_brief_list()

            # --- lihat brief ---
            if low in ("lihat brief", "show brief", "/brief", "brief sekarang"):
                return self._fmt_brief(target or self.default_target)

            # --- manifest ---
            if low in ("manifest", "generate manifest", "/manifest"):
                return self._fmt_manifest(target or self.default_target)

            # --- hapus brief ---
            m = re.match(r'^(?:hapus|clear|delete|reset)\s+brief(?:\s+(\S+))?$', low)
            if m:
                t = m.group(1) or target or self.default_target
                if self.brief_store.clear_brief(t):
                    return f"🗑 Brief '{t}' dihapus."
                return f"ℹ️ Tidak ada brief untuk '{t}'."

                        # --- ingest file (CSV/JSON/XLSX/PDF) -> brief ---
            # Bentuk: "scopes.csv" | "buka scopes.csv" | "import scopes.csv untuk shopify.com"
            fname_token = None
            stripped = low.strip()
            file_exts = ('.csv', '.json', '.xlsx', '.xls', '.pdf')
            if stripped.endswith(file_exts) and ' ' not in stripped:
                fname_token = stripped
            else:
                for kw in ('buka file', 'from file', 'buka', 'baca', 'upload',
                           'import', 'file', 'load', 'read', 'muat'):
                    if (len(stripped) > len(kw) and stripped.startswith(kw)
                            and stripped[len(kw)] == ' '):
                        fname_token = stripped[len(kw):].strip()
                        break
            if fname_token:
                rest = fname_token
                target_override = None
                path_part = rest
                for sep in (' untuk ', ' dari ', ' target ', ' pada '):
                    idx = rest.find(sep)
                    if idx != -1:
                        target_override = rest[idx + len(sep):].strip()
                        path_part = rest[:idx].strip()
                        break
                path_part = path_part.strip(chr(34)).strip(chr(39)).strip()
                if not path_part or not path_part.lower().endswith(file_exts):
                    fname_token = None  # bukan perintah file -> biarkan handler lain
            if fname_token:
                if not path_part:
                    return "ℹ️ Sebutkan path berkas, mis. 'buka scopes.csv' atau 'scopes.csv'."
                path = self._resolve_file(path_part)
                if not path:
                    return (f"⚠️ Berkas tidak ditemukan: {path_part}\n"
                            "📂 Sistem cari di: folder kerja, folder "
                            "'DIALOGIC_COPILLOT'/'data/file', dan '~/.arc'.")
                return self.ingest_file(path, target=target_override or target)

            # --- tambah scope ---
            m = re.match(r'^(?:tambah|add|set|update)\s+(?:in\s+)?scope\s+(.+)$', message, re.I)
            if m and target:
                items = self.brief_store.add_item(target, "scope", m.group(1))
                return (f"✅ Scope diperbarui untuk {target} ({len(items['scope'])} item):\n"
                        + "\n".join(f"  • {it}" for it in items["scope"]))

            # --- tambah out of scope ---
            m = re.match(r'^(?:tambah|add|set)\s+(?:out\s*[- ]?of\s*scope|oos)\s+(.+)$',
                         message, re.I)
            if m and target:
                items = self.brief_store.add_item(target, "out_of_scope", m.group(1))
                return (f"✅ Out-of-scope diperbarui untuk {target} "
                        f"({len(items['out_of_scope'])} item):\n"
                        + "\n".join(f"  🚫 {it}" for it in items["out_of_scope"]))

            # --- tambah requirement ---
            m = re.match(r'^(?:tambah|add|set)\s+(?:requirement|requirements|persyaratan)\s+(.+)$',
                         message, re.I)
            if m and target:
                items = self.brief_store.add_item(target, "requirements", m.group(1))
                return (f"✅ Requirements diperbarui untuk {target} "
                        f"({len(items['requirements'])} item):\n"
                        + "\n".join(f"  📌 {it}" for it in items["requirements"]))

            # --- set rules ---
            m = re.match(r'^(?:set|update)\s+rules?\s+(.+)$', message, re.I)
            if m and target:
                return self._set_rules(target, m.group(1))

            # --- intel Google VRP ---
            if low in ("program google", "google programs", "daftar program google",
                       "program google vrp", "list google programs"):
                return self._google_programs_summary()

            m = re.match(r'^(?:cek|check|test)\s+scope\s+(\S+)$', low)
            if m:
                return self._google_scope_check(m.group(1))

            # --- pertanyaan: "apa scope target X?" / "apa out of scope X?" ---
            q = re.match(
                r'^(?:apa|what|sebutkan|list|tampilkan|show|lihat)\s+'
                r'(scope|out\s+of\s+scope|rules|requirements)'
                                r'(?:\s+(?:untuk|dari|target)\s+)?\s*([^\s?]*)\s*\??$', low)
            if q:
                field = q.group(1).lower().replace(" ", "_")
                t = (q.group(2) or "").strip().lower() or target or self.default_target
                return self._fmt_field(t, field)

            # --- pertanyaan singkat: "rules untuk X" / "scope target X" ---
            q2 = re.match(
                r'^(scope|out\s+of\s+scope|rules|requirements)\s+'
                                r'(?:untuk|dari|target)\s+([^\s?]+)\s*\??$', low)
            if q2:
                field = q2.group(1).lower().replace(" ", "_")
                t = q2.group(2).strip().lower() or target or self.default_target
                return self._fmt_field(t, field)

            # --- "target X" (singkat) -> pindah target ---
            m = re.match(r'^target\s+(\S+)$', low)
            if m:
                t = m.group(1).lower()
                self.conversation.start_conversation(target_domain=t)
                return f"🎯 Target diganti: {t}\nKirim detail bounty atau /bantuan."

            # --- intent membuat brief: "buat brief untuk X, scope ..." ---
            if re.search(r'\b(buat|bikin|catat|simpan|set|update|isi)\s+brief\b', low) \
               and re.search(r'\b(scope|rules|requirements|target|deskripsi)\b', low):
                return self._parse_and_save_brief(message, target)

            # --- deteksi paste brief bounty / format natural ---
            if self._is_brief_text(message):
                return self._parse_and_save_brief(message, target)

            # --- bentuk bebas: "<field> <items>" tanpa kata sandang ---
            # Contoh: "scope app.a.com, *.a.com" | "out of scope staging.x"
            #         "rules bounty: $500; response time: 72 jam"
            #         "requirements repro wajib, bukti PoC"
            # Diletakkan SETELAH _is_brief_text agar kalimat brief panjang
            # (multi-field, len>40, ≥2 keyword) tetap ke _parse_and_save_brief,
            # dan SETELAH pertanyaan q/q2 agar "field untuk X?" tetap jadi pertanyaan.
            mfb = re.match(
                r'^(scope|in\s+scope|in-scope|out\s+of\s+scope|out-of-scope|'
                r'out_of_scope|deskripsi|description|requirements?|'
                r'(?:aturan|peraturan|ketentuan|rules?))\s+'
                r'(?!untuk\b|dari\b|target\b)(\S.*?)\s*$', message, re.I)
            if mfb and target:
                return self._apply_bare_field(target, mfb.group(1), mfb.group(2))

            return None
        except Exception as e:
            return f"⚠️ Brief handler error (safe): {e}"

        # ==================================================================
    # INGEST FILE (CSV / JSON / XLSX / PDF -> ProgramBrief)
    # ==================================================================
    def _resolve_file(self, fname: str) -> Optional[str]:
        """Selesaikan path relatif / nama berkas ke path absolut yang ada."""
        if not fname:
            return None
        fn = os.path.expanduser(fname.strip().strip(chr(34)).strip(chr(39)).strip())
        if os.path.isfile(fn):
            return os.path.abspath(fn)
        if os.path.isabs(fn):
            return None
        exts = ['', '.csv', '.json', '.xlsx', '.xls', '.pdf']
        lower = fn.lower()
        roots = [
            os.getcwd(),
            os.path.dirname(os.path.abspath(__file__)),
            os.path.join(os.getcwd(), 'DIALOGIC_COPILLOT'),
            os.path.join(os.getcwd(), 'data', 'file'),
            os.path.join(os.path.expanduser('~'), '.arc'),
        ]
        for root in roots:
            if not os.path.isdir(root):
                continue
            for ext in exts:
                cand = os.path.join(root, fn if lower.endswith(ext) else fn + ext)
                if os.path.isfile(cand):
                    return os.path.abspath(cand)
        return None

    @staticmethod
    def _truthy(value: Any) -> bool:
        s = str(value or '').strip().lower()
        return s in ('true', 'yes', '1', 'eligible', 'ya', 'benar')

    def _detect_columns(self, headers: List[str]) -> Dict[str, Optional[str]]:
        """Petaan header berkas ke kolom kanonik (nama kolom bebas, case-insensitive)."""
        norm = {}
        for h in headers:
            key = (h or '').strip().lower().replace(' ', '_')
            if key and key not in norm:
                norm[key] = h

        def find(candidates):
            for c in candidates:
                if c in norm:
                    return norm[c]
            return None

        return {
            'identifier': find(['identifier', 'asset', 'target', 'domain', 'host',
                                'hostname', 'ip', 'ip_address', 'endpoint',
                                'in_scope', 'target_identifier']),
            'out_of_scope': find(['out_of_scope', 'out-of-scope', 'excluded',
                                  'not_in_scope', 'oos']),
            'description': find(['description', 'deskripsi', 'notes', 'instruction',
                                 'instruction_text', 'about', 'summary', 'details']),
            'max_severity': find(['max_severity', 'severity', 'bounty_tier',
                                  'highest_reward', 'max_reward']),
            'elig_bounty': find(['eligible_for_bounty', 'bounty', 'reward_eligible',
                                 'eligible_bounty', 'rewardable']),
            'elig_sub': find(['eligible_for_submission', 'submission_eligible',
                              'eligible_submission', 'eligible_report']),
            'asset_type': find(['asset_type', 'type', 'asset', 'asset_classification']),
            'environment': find(['environment', 'env', 'environment_type']),
            'requirements': find(['requirements', 'persyaratan', 'requirements_text']),
        }

    def _read_rows(self, path: str) -> List[Dict[str, Any]]:
        ext = os.path.splitext(path)[1].lower()
        if ext == '.json':
            return self._read_json(path)
        if ext in ('.xlsx', '.xls'):
            return self._read_excel(path)
        if ext == '.pdf':
            return self._read_pdf(path)
        if ext in ('.docx', '.doc'):
            return self._read_docx(path)
        return self._read_csv(path)

    def _read_csv(self, path: str) -> List[Dict[str, str]]:
        import csv as _csv
        rows = []
        with open(path, 'r', encoding='utf-8-sig', newline='') as f:
            sample = f.read(8192)
            f.seek(0)
            try:
                delim = _csv.Sniffer().sniff(sample, delimiters=',;').delimiter
            except Exception:
                delim = ','
            for row in _csv.DictReader(f, delimiter=delim):
                rows.append({(k or '').strip(): (v if v is not None else '')
                             for k, v in row.items()})
        return rows

    def _read_json(self, path: str) -> List[Dict[str, Any]]:
        with open(path, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
        if isinstance(data, list):
            return [r for r in data if isinstance(r, dict)]
        if isinstance(data, dict):
            for key in ('scope', 'scopes', 'entries', 'items', 'targets',
                        'programs', 'assets', 'findings', 'data', 'results'):
                sub = data.get(key)
                if isinstance(sub, list) and sub and isinstance(sub[0], dict):
                    return sub
                if isinstance(sub, dict):
                    for k2 in ('scope', 'scopes', 'entries', 'items', 'targets', 'assets'):
                        s2 = sub.get(k2)
                        if isinstance(s2, list) and s2 and isinstance(s2[0], dict):
                            return s2
            return [data]
        return []

    def _read_excel(self, path: str) -> List[Dict[str, Any]]:
        try:
            import openpyxl
        except Exception:
            return []
        rows = []
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        header = None
        for ridx, row in enumerate(ws.iter_rows(values_only=True)):
            if ridx == 0:
                header = [str(c).strip() if c is not None else '' for c in (row or [])]
                continue
            if not header:
                continue
            vals = list(row or [])
            rows.append({header[i] if i < len(header) else f'col{i}': (v if v is not None else '')
                         for i, v in enumerate(vals)})
        return rows

    def _read_pdf(self, path: str) -> List[Dict[str, Any]]:
        rows = []
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for table in (page.extract_tables() or []):
                        header = None
                        for i, r in enumerate(table):
                            if i == 0:
                                header = [str(c).strip() if c else '' for c in r]
                                continue
                            if header:
                                rows.append({header[j] if j < len(header) else f'col{j}':
                                             (r[j] if j < len(r) else '')
                                             for j in range(len(header))})
        except Exception:
            pass
        if rows:
            return rows
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            out = []
            for p in reader.pages:
                out.extend((p.extract_text() or '').splitlines())
            return [{'field': ln} for ln in out if ln.strip()]
        except Exception:
            return []

    def _read_docx(self, path: str) -> List[Dict[str, Any]]:
        rows = []
        try:
            from docx import Document
            doc = Document(path)
            # Extract table data if available
            for table in doc.tables:
                if not table.rows:
                    continue
                # Use first row as header
                header = [cell.text.strip() for cell in table.rows[0].cells]
                if not header or all(not h for h in header):
                    continue
                for row in table.rows[1:]:
                    values = [cell.text.strip() for cell in row.cells]
                    row_dict = {}
                    for i, h in enumerate(header):
                        if i < len(values):
                            row_dict[h] = values[i]
                    if any(v for v in row_dict.values()):
                        rows.append(row_dict)
            # If no tables, extract text from paragraphs
            if not rows:
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        rows.append({'field': text})
        except Exception:
            pass
        return rows

    def _rows_to_brief(self, rows: List[Dict[str, Any]], target: str,
                       platform: str, source_label: str) -> Optional[Dict[str, Any]]:
        """Klasifikasikan baris ekspor scope -> struktur ProgramBrief."""
        if not rows:
            return None
        cols = self._detect_columns(list(rows[0].keys()))
        ident_col = cols['identifier']
        desc_col = cols['description']
        oos_col = cols['out_of_scope']
        sev_col = cols['max_severity']
        bounty_col = cols['elig_bounty']
        sub_col = cols['elig_sub']
        atype_col = cols['asset_type']
        req_col = cols['requirements']
        sev_rank = {'none': 0, 'low': 1, 'medium': 2, 'high': 3, 'critical': 4, 'n/a': 0, '': 0}
        in_scope, out_scope, asset_types, reqs, notes, accepted = [], [], set(), set(), [], set()
        max_sev = 'none'
        has_bounty = False
        for r in rows:
            if ident_col:
                idv = _clean(str(r.get(ident_col)))
            else:
                idv = next((_clean(str(r[k])) for k in r
                            if r.get(k) and str(r.get(k)).strip()), '')
            if not idv:
                continue
            instr = _clean(str(r.get(desc_col) if desc_col else ''))
            atype = _clean(str(r.get(atype_col) if atype_col else '')).lower()
            if atype:
                asset_types.add(atype)
            sev = _clean(str(r.get(sev_col) if sev_col else '')).lower() or 'none'
            if sev_rank.get(sev, 0) > sev_rank.get(max_sev, 0):
                max_sev = sev
            bounty = self._truthy(r.get(bounty_col) if bounty_col else False)
            sub = self._truthy(r.get(sub_col) if sub_col else False)
            explicit_oos = bool(oos_col and self._truthy(r.get(oos_col)))
            is_oos = explicit_oos or not (bounty or sub)
            if is_oos:
                out_scope.append(idv)
            else:
                in_scope.append(idv)
                has_bounty = has_bounty or bounty
            if sev in sev_rank and sev_rank[sev] > 0:
                accepted.add(sev)
            if req_col and r.get(req_col):
                for it in _split_items(str(r.get(req_col))):
                    if it:
                        reqs.add(it)
            if instr:
                notes.append(f"{idv}: {instr}")
        if not in_scope and not out_scope:
            return None
        rules = {
            'bounty': 'eligible (ada reward)' if has_bounty else 'tidak ada reward yang tercantum',
            'max_severity': max_sev,
            'response_time': 'tidak ditentukan',
            'asset_types': ', '.join(sorted(asset_types)) if asset_types else 'mixed',
            'source': source_label,
        }
        if accepted:
            rules['severity_tiers'] = ', '.join(
                sorted(accepted, key=lambda s: sev_rank.get(s, 0), reverse=True))
        if notes:
            rules['_asset_notes'] = notes[:50]
        return {
            'program_name': target,
            'platform': platform,
            'target_domain': target,
            'description': (f"Brief diimpor dari {source_label} "
                            f"({len(rows)} asset: {len(in_scope)} in-scope, "
                            f"{len(out_scope)} out-of-scope)."),
            'scope': in_scope,
            'out_of_scope': out_scope,
            'rules': rules,
            'requirements': sorted(reqs),
            'allowed_operations': ['recon', 'scan', 'exploit'],
            'source': source_label,
            'updated_at': datetime.now().isoformat(),
        }

    def ingest_file(self, path: str, target: Optional[str] = None) -> str:
        """Baca berkas CSV/JSON/XLSX/PDF dan simpan/merge ke brief target."""
        try:
            ext = os.path.splitext(path)[1].lower()
            if ext not in ('.csv', '.json', '.xlsx', '.xls', '.pdf'):
                return (f"⚠️ Format {ext or '(tanpa ekstensi)'} belum didukung. "
                        "ARC mendukung: csv, json, xlsx, pdf.")
            rows = self._read_rows(path)
            if not rows:
                return (f"⚠️ Tidak berhasil baca data dari {os.path.basename(path)} "
                        "(tabel/header tidak terdeteksi). Pastikan berkas berisi tabel.")
            conv = self.conversation.current_conversation
            tgt = (target or (conv or {}).get('target_domain')
                   or self.default_target)
            keys = str(list(rows[0].keys()))
            is_h1 = ('eligible_for' in keys
                     or 'hackerone' in os.path.basename(path).lower())
            platform = 'hackerone' if is_h1 else 'manual_import'
            brief = self._rows_to_brief(rows, tgt, platform,
                                        f"file:{os.path.basename(path)}")
            if not brief:
                return (f"⚠️ Berkas dibaca ({os.path.basename(path)}) tapi tidak ada "
                        "asset yang terklasifikasi (kolom identifier tidak ditemukan).")
            merged = self.brief_store.upsert(tgt, brief)
            ms = merged.get('rules', {}).get('max_severity', '')
            ub = merged.get('rules', {}).get('bounty', '')
            lines = [
                f"📁 Import dari '{os.path.basename(path)}' -> brief '{tgt}'",
                f"   in-scope : {len(merged.get('scope', []))}  | out-of-scope : {len(merged.get('out_of_scope', []))}  | rules: {len(merged.get('rules', {}))}",
                f"   max_severity: {ms}  | bounty: {ub}",
                f"   📦 {self.brief_store._path(tgt)}",
                "   Ketik /manifest untuk generate manifest otorisasi, atau 'lihat brief'.",
            ]
            return '\n'.join(lines)
        except Exception as e:
            return f"⚠️ Ingest file error (safe): {e}"

    # ==================================================================
    # OPERASI BRIEF
    # ==================================================================
    def _set_rules(self, target: str, rules_text: str) -> str:
        """Parse 'k1: v1; k2: v2' -> merge ke brief target."""
        rules: Dict[str, Any] = {}
        notes = []
        for chunk in re.split(r'[;]', rules_text):
            chunk = chunk.strip()
            if not chunk:
                continue
            if ":" in chunk:
                k, _, v = chunk.partition(":")
                k = k.strip()
                v = v.strip()
                if k:
                    rules[k] = v
            else:
                notes.append(chunk)
        if notes:
            rules["_notes"] = notes
        if not rules:
            return "ℹ️ Format rules: 'set rules kunci: nilai; kunci2: nilai2'"
        merged = self.brief_store.upsert(target, {"rules": rules})
        lines = [f"✅ Rules diperbarui untuk {target}:"]
        for k, v in merged.get("rules", {}).items():
            if k == "_notes":
                lines.append(f"  📝 notes: {v}")
            else:
                lines.append(f"  • {k}: {v}")
        return "\n".join(lines)

    def _is_brief_text(self, text: str) -> bool:
        low = text.lower()
        # (a) format colon / paste dashboard: "Scope: ..."
        markers = ["scope:", "out of scope:", "out-of-scope:", "deskripsi:",
                   "rules:", "requirements:"]
        hits = sum(1 for mk in markers if mk in low)
        if hits >= 2:
            return True
        if hits >= 1 and (len(text) > 120 or "program" in low
                          or "bounty" in low or "brief" in low):
            return True
        # (b) format natural tanpa titik dua:
        #     "target X scope Y out of scope Z rules W requirements V"
        #     Hitung kata kunci tak tumpang-tindih (frasa panjang lebih dulu),
        #     supaya "scope" di dalam "out of scope" tidak dihitung ganda.
        if len(text) > 40:
            natural_keys = ["deskripsi", "out of scope", "scope", "rules",
                            "requirements"]
            nhits = 0
            occupied = [False] * (len(low) + 1)
            for kw in sorted(natural_keys, key=len, reverse=True):
                for mm in re.finditer(r'\b' + re.escape(kw) + r'\b', low):
                    st, e = mm.span()
                    if occupied[st] or occupied[e - 1]:
                        continue
                    for k in range(st, e):
                        occupied[k] = True
                    nhits += 1
            if nhits >= 2:
                return True
        # (c) intent eksplisit: "buat brief untuk X, scope ..."
        if "brief" in low and re.search(r'\b(scope|rules|requirements|target|deskripsi)\b', low):
            return True
        return False

    @staticmethod
    def _has_colon_sections(text: str) -> bool:
        """True bila teks pakai format 'key: value' (bukan perlu dinormalisasi)."""
        low = text.lower()
        return any(mk in low for mk in [
            "scope:", "out of scope:", "out-of-scope:",
            "deskripsi:", "rules:", "requirements:"])

    @staticmethod
    def _to_colon_format(text: str) -> str:
        """Ubah format natural 'target X scope Y rules Z' menjadi 'key: value'."""
        pattern = re.compile(
            r'\b(deskripsi|target|scope|out\s+of\s+scope|requirements?|rules?)\b',
            re.IGNORECASE)
        matches = list(pattern.finditer(text))
        if len(matches) < 2:  # bukan kalimat brief natural
            return text
        lines = []
        for i, m in enumerate(matches):
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            seg = text[start:end].strip(" ,;:-")
            if not seg:
                continue
            # kata hubung -> koma agar _split_items bisa memecah item
            seg = re.sub(r'\s+(?:dan|and|serta|atau)\s+', ', ', seg, flags=re.IGNORECASE)
            seg = re.sub(r'\s*\d+\)\s*', '; ', seg)
            key = m.group(1).lower().strip()
            if key.startswith("out"):
                key = "out of scope"
            elif key in ("rule", "rules"):
                key = "rules"
            elif key in ("requirement", "requirements"):
                key = "requirements"
            lines.append(f"{key}: {seg}")
        return "\n".join(lines)

    def _parse_and_save_brief(self, text: str, current_target: Optional[str]) -> str:
        # Normalisasi format natural ("target X scope Y rules Z") -> 'key: value' bila perlu
        parse_text = text if self._has_colon_sections(text) else self._to_colon_format(text)
        target = (self._extract_target(parse_text) or self._extract_target(text)
                  or current_target or self.default_target)
        program_name = self._extract_program_name(parse_text) or self._extract_program_name(text)
        brief = parse_program_brief(
            parse_text,
            platform=self._extract_platform(parse_text) or "manual_dialog",
            program_name=program_name,
            target_domain=target,
        )
        # Beri tahu kalau teks tidak mengandung marker section sama sekali
        if not (brief["scope"] or brief["out_of_scope"] or brief["rules"]
                or brief["requirements"] or brief["description"]):
            return ("ℹ️ Tidak terdeteksi section brief (scope/deskripsi/rules/requirements).\n"
                    "Contoh format:\n"
                    "  Deskripsi: ...\n  Scope: app.example.com, *.example.com\n"
                    "  Out of scope: staging.example.com\n  Rules: ...\n  Requirements: ...")
        path = self.brief_store.save_brief(brief)
        return (f"✅ Brief untuk '{target}' tersimpan!\n"
                f"📦 File: {path}\n"
                f"  Deskripsi    : {len(brief['description'])} karakter\n"
                f"  Scope        : {len(brief['scope'])} item\n"
                f"  Out of scope : {len(brief['out_of_scope'])} item\n"
                f"  Rules        : {len(brief['rules'])} entri\n"
                f"  Requirements : {len(brief['requirements'])} item\n"
                f"Ketik /manifest untuk generate manifest otorisasi.")

    def _apply_bare_field(self, target: str, field_token: str,
                          payload: str) -> str:
        """Tuliskan/append satu field brief lewat bahasa alami bebas
        (tanpa kata sandang 'tambah/set').

        Dipanggil oleh handle_brief_command untuk bentuk:
            scope app.a.com, *.a.com
            out of scope staging.x
            rules bounty: $500; response time: 72 jam
            requirements repro wajib, bukti PoC
            deskripsi ini adalah bounty bank
        """
        ft = field_token.lower().replace("-", " ").replace("_", " ")
        if ft.startswith("out") or "out of" in ft or ft == "excluded":
            field = "out_of_scope"
        elif ft.startswith("in scope") or ft == "scope":
            field = "scope"
        elif ft in ("deskripsi", "description"):
            field = "description"
        elif ft.startswith("aturan") or ft.startswith("peraturan") or \
                ft.startswith("ketentuan") or ft.startswith("rule"):
            field = "rules"
        else:
            field = "requirements"

        payload = (payload or "").strip()
        if not payload:
            return "ℹ️ Beri itemnya, mis. 'scope app.a.com, *.a.com'."

        if field == "rules":
            return self._set_rules(target, payload)
        if field == "description":
            brief = self.brief_store.upsert(target, {"description": payload})
            return (f"✅ Deskripsi disimpan untuk {target} "
                    f"({len(brief.get('description', ''))} karakter).")

        items = self.brief_store.add_item(target, field, payload)
        got = items.get(field, []) or []
        label = field.replace("_", " ")
        return (f"✅ {label.title()} tersimpan untuk {target} "
                f"({len(got)} item):\n"
                + "\n".join(f"  • {it}" for it in got))

    def _extract_program_name(self, text: str) -> Optional[str]:
        m = re.search(r'(?:program|program_name)\s*[:=]\s*([^\n\r;]{2,60})', text, re.IGNORECASE)
        return m.group(1).strip() if m else None

    @staticmethod
    def _extract_platform(text: str) -> Optional[str]:
        m = re.search(r'(?:platform)\s*[:=]\s*([a-zA-Z0-9 _-]+)', text, re.IGNORECASE)
        return m.group(1).strip() if m else None

    # ==================================================================
    # FORMATTER OUTPUT
    # ==================================================================
    def _fmt_brief_list(self) -> str:
        briefs = self.brief_store.list_briefs()
        if not briefs:
            return "📭 Belum ada brief tersimpan. Mulai dengan paste teks bounty."
        lines = [f"📚 Brief tersimpan ({len(briefs)}):"]
        for b in briefs:
            lines.append(f"  • {b}")
        return "\n".join(lines)

    def _fmt_brief(self, target: str) -> str:
        if not target:
            return "Belum ada target aktif. Mulai dengan '/mulai <target>'."
        brief = self.brief_store.get_brief(target)
        if not brief or not (brief.get("scope") or brief.get("rules")
                             or brief.get("requirements") or brief.get("description")):
            return (f"ℹ️ Belum ada brief untuk '{target}'.\n"
                    f"Paste teks bounty atau perintah: tambah scope <item>; set rules k: v")
        lines = [f"📋 BRIEF — {target}"]
        if brief.get("program_name"):
            lines.append(f"  Program  : {brief['program_name']}")
        if brief.get("platform"):
            lines.append(f"  Platform : {brief['platform']}")
        if brief.get("description"):
            lines.append(f"  Deskripsi: {brief['description'][:220]}")
        scope = brief.get("scope", [])
        lines.append(f"  🛡 Scope ({len(scope)}):")
        lines += [f"    • {s}" for s in scope] or ["    (kosong)"]
        oos = brief.get("out_of_scope", [])
        lines.append(f"  🚫 Out of scope ({len(oos)}):")
        lines += [f"    • {s}" for s in oos] or ["    (kosong)"]
        rules = brief.get("rules", {})
        if rules:
            lines.append(f"  📜 Rules ({len(rules)}):")
            for k, v in rules.items():
                if k == "_notes":
                    lines.append(f"    📝 {v}")
                else:
                    lines.append(f"    • {k}: {v}")
        reqs = brief.get("requirements", [])
        lines.append(f"  📌 Requirements ({len(reqs)}):")
        lines += [f"    • {r}" for r in reqs] or ["    (kosong)"]
        return "\n".join(lines)

    def _fmt_field(self, target: str, field: str) -> str:
        """Tampilkan satu bagian brief saja (scope/rules/requirements/out_of_scope)."""
        brief = self.brief_store.get_brief(target)
        if not brief:
            return f"ℹ️ Belum ada brief untuk '{target}'."
        if field in ("scope", "out_of_scope", "requirements"):
            items = brief.get(field, []) or []
            label = field.replace("_", " ")
            if not items:
                return f"ℹ️ {label} untuk '{target}' masih kosong."
            return f"📋 {label.title()} — {target}:\n" + "\n".join(f"  • {i}" for i in items)
        if field == "rules":
            rules = brief.get("rules", {}) or {}
            if not rules:
                return f"ℹ️ rules untuk '{target}' masih kosong."
            lines = [f"📋 Rules — {target}:"]
            for k, v in rules.items():
                if k == "_notes":
                    lines.append(f"  📝 {v}")
                else:
                    lines.append(f"  • {k}: {v}")
            return "\n".join(lines)
        return self._fmt_brief(target)

    def _fmt_manifest(self, target: str) -> str:
        if not target:
            return "Belum ada target aktif. Mulai dengan '/mulai <target>'."
        brief = self.brief_store.get_brief(target)
        if not brief or not brief.get("scope"):
            return f"ℹ️ Belum ada scope untuk '{target}'. Tambahkan dulu: 'tambah scope <item>'"
        try:
            manifest = brief_to_manifest(brief)
            return ("📜 MANIFEST (untuk ScopeSovereigntyGuard):\n"
                    + json.dumps(manifest, indent=2, ensure_ascii=False))
        except Exception as e:
            return f"⚠️ Gagal generate manifest: {e}"

    # ==================================================================
    # INTEL GOOGLE VRP (lazy import, opsional)
    # ==================================================================
    def _google(self) -> Optional[Any]:
        if not self.enable_google_vrp:
            return None
        if self._google_integrator is not None:
            return self._google_integrator
        try:
            from BROWSER_SECURITY_RESEARCH.google_vrp_integrator import GoogleVRPIntegrator
            self._google_integrator = GoogleVRPIntegrator()
        except Exception as e:
            self._google_error = str(e)
            self._google_integrator = False
        return self._google_integrator if self._google_integrator else None

    def _google_programs_summary(self) -> str:
        g = self._google()
        if g is None:
            err = f" ({self._google_error})" if self._google_error else ""
            return f"⚠️ Google VRP integrator tidak tersedia{err}."
        try:
            progs = g.get_all_google_programs()
            lines = [f"⚡ Google Vulnerability Reward Programs ({len(progs)}):"]
            for key, p in sorted(progs.items()):
                status = p.get("status", "")
                lines.append(f"  • {key}: {p.get('name')} — {p.get('bounty_range')}"
                             + (f" [{status}]" if status else ""))
            return "\n".join(lines)
        except Exception as e:
            return f"⚠️ Gagal mengambil program Google: {e}"

    def _google_scope_check(self, url: str) -> str:
        g = self._google()
        if g is None:
            err = f" ({self._google_error})" if self._google_error else ""
            return f"⚠️ Google VRP integrator tidak tersedia{err}."
        try:
            in_scope = g.is_target_in_google_scope(url)
            matches = g.find_matching_google_programs(url)
            head = "🟢 IN SCOPE" if in_scope else "🔴 NOT IN SCOPE"
            lines = [f"{head} → {url}"]
            if matches:
                lines.append("✅ Program yang cocok:")
                for key in matches:
                    p = g.google_programs.get(key, {})
                    lines.append(f"   • {key} — {p.get('name')} ({p.get('bounty_range')})")
            else:
                lines.append("ℹ️ Tidak cocok dengan pola publik program Google.")
                lines.append("💡 Verifikasi manual: https://bughunters.google.com/report/new")
            return "\n".join(lines)
        except Exception as e:
            return f"⚠️ Gagal cek scope Google: {e}"

    # ==================================================================
    # FALLBACK & BANTUAN
    # ==================================================================
    def _fallback_reply(self, message: str) -> str:
        low = message.lower()
        greetings = ("halo", "hai", "hi", "hey", "hello", "pagi", "siang",
                     "sore", "malam", "assalamualaikum", "assalamu'alaikum")
        if any(g in low for g in greetings):
            return ("Halo! 👋 Saya ARC v7.6.\n"
                    "Bilang saja langsung, misalnya:\n"
                    "  'target bank.xyz, scope *.bank.xyz, api.bank.xyz, "
                    "rules xss: allowed, out of scope staging.bank.xyz'\n"
                    "Atau ketik /bantuan.")
        try:
            interp = self.commands.interpret_command(message)
        except Exception:
            interp = {}
        action = interp.get("interpreted_action")
        if action and action != "unknown":
            target = interp.get("parameters", {}).get("target") or "?"
            return (f"🛠 Perintah terdeteksi: {action} (target: {target})\n"
                    "⚠️ Eksekusi butuh sesi ARC penuh (python arc_main.py) + "
                    "approval human-in-the-loop.\n"
                    "📌 Untuk sekarang: beri tahu saya detail bounty target ini, "
                    "nanti semua operasi dibatasi otomatis oleh ScopeSovereigntyGuard.")
        return ("🤖 Saya ARC v7.6 — siap tukar informasi bounty.\n"
                "Coba:\n"
                "  • /mulai <target>          — mulai dialog untuk target apa saja\n"
                "  • paste teks brief bounty  — saya parse & simpan otomatis\n"
                "  • tambah scope <item>      — tambah in-scope bertahap\n"
                "  • /lihat brief, /manifest, /daftar brief, /bantuan")

    @property
    def help_text(self) -> str:
        return (
            "📖 BANTUAN ARC CHAT v7.6\n"
            "─" * 40 + "\n"
            "🎯 TARGET (bebas: domain, IP, program, nama apa saja)\n"
            "  /mulai <target> | mulai target <target> | ganti target <target>\n"
            "  target\n"
            "\n"
            "📋 BRIEF BOUNTY\n"
            "  Paste teks bounty (Deskripsi/Scope/Rules/Requirements) → parse otomatis\n"
            "  Bisa juga natural: target A, scope *.a.com, rules xss: allowed; out of scope B\n"
            "  tambah scope <item> | tambah out of scope <item>\n"
            "  tambah requirement <item> | set rules kunci: nilai; k2: v2\n"
            "  lihat brief | daftar brief | hapus brief <target>\n"
            "\n"
            "🛡 MANIFEST\n"
            "  manifest — generate manifest otorisasi (ScopeSovereigntyGuard)\n"
            "\n"
            "⚡ INTEL GOOGLE VRP\n"
            "  program google | cek scope <url>\n"
            "\n"
            "💬 Lainnya\n"
            "  /bantuan"
        )


if __name__ == "__main__":
    import sys
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass
    _engine = ArcChatEngine()
    print("=== ARC Chat Engine (diagnostik) ===")
    print(_engine.handle("/bantuan"))
