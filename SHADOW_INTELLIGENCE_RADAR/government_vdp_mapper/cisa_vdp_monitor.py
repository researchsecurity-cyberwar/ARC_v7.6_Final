import requests
from bs4 import BeautifulSoup
import re
import tempfile
import os

class CISAVDPMonitor:
    """
    US federal VDP tracking.
    Memantau program VDP federal AS melalui situs CISA dengan pendekatan dinamis.
    Mendukung ekstraksi konten dari dokumen PDF, DOCX, dan XLSX secara otomatis.
    """
    
    def __init__(self):
        # Gunakan homepage yang stabil, bukan URL spesifik yang bisa berubah
        self.cisa_home_url = "https://www.cisa.gov"
    
    def get_federal_vdp_programs(self):
        """Dapatkan program VDP federal dari sumber CISA yang valid."""
        try:
            # Mulai dari homepage CISA yang pasti aktif
            response = requests.get(self.cisa_home_url, timeout=10)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Cari link yang relevan dengan VDP/reporting
                vdp_links = []
                for link in soup.find_all('a', href=True):
                    link_text = link.get_text().strip().lower()
                    href = link['href']
                    
                    # Cari kata kunci yang menunjukkan halaman VDP/reporting
                    if any(keyword in link_text for keyword in [
                        'vulnerability', 'report', 'security', 'disclose', 
                        'contact', 'submit', 'incident'
                    ]):
                        # Normalisasi URL menjadi absolut
                        if href.startswith('/'):
                            full_url = f"{self.cisa_home_url}{href}"
                        elif href.startswith('http'):
                            full_url = href
                        else:
                            continue
                        
                        # Validasi apakah URL benar-benar aktif
                        if self._is_url_active(full_url):
                            program_info = {
                                'agency': 'CISA',
                                'vdp_url': full_url,
                                'country': 'usa',
                                'type': 'federal'
                            }
                            
                            # Jika ini file dokumen, ekstrak kontennya
                            if self._is_document_url(full_url):
                                document_content = self._download_and_extract_document(full_url)
                                if document_content:
                                    program_info['document_content'] = document_content
                            
                            vdp_links.append(program_info)
                
                return vdp_links
                    
        except Exception as e:
            print(f"⚠️ Failed to fetch CISA programs: {e}")
        
        return []
    
    def _is_url_active(self, url):
        """Validasi apakah URL mengembalikan status 200 OK."""
        try:
            response = requests.head(url, timeout=5, allow_redirects=True)
            return response.status_code == 200
        except:
            return False
    
    def _is_document_url(self, url):
        """Cek apakah URL mengarah ke file dokumen yang bisa diparse."""
        doc_extensions = ['.pdf', '.docx', '.xlsx', '.xls']
        return any(url.lower().endswith(ext) for ext in doc_extensions)
    
    def _download_and_extract_document(self, url):
        """Download dan ekstrak konten dari dokumen (PDF/DOCX/XLSX)."""
        try:
            # Download file ke temporary location
            response = requests.get(url, timeout=30)
            if response.status_code != 200:
                return None
            
            # Tentukan ekstensi file
            file_ext = self._get_file_extension(url)
            if not file_ext:
                return None
            
            # Simpan ke file temporary
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp_file:
                tmp_file.write(response.content)
                tmp_path = tmp_file.name
            
            try:
                # Ekstrak konten berdasarkan jenis file
                if file_ext == '.pdf':
                    return self._extract_pdf_content(tmp_path)
                elif file_ext == '.docx':
                    return self._extract_docx_content(tmp_path)
                elif file_ext in ['.xlsx', '.xls']:
                    return self._extract_xlsx_content(tmp_path)
                else:
                    return None
                    
            finally:
                # Hapus file temporary setelah selesai
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    
        except Exception as e:
            print(f"⚠️ Failed to process document {url}: {e}")
            return None
    
    def _get_file_extension(self, url):
        """Ekstrak ekstensi file dari URL."""
        import urllib.parse
        parsed_url = urllib.parse.urlparse(url)
        path = parsed_url.path.lower()
        
        if path.endswith('.pdf'):
            return '.pdf'
        elif path.endswith('.docx'):
            return '.docx'
        elif path.endswith(('.xlsx', '.xls')):
            return '.xlsx'
        else:
            return None
    
    def _extract_pdf_content(self, pdf_path):
        """Ekstrak konten dari file PDF menggunakan pdfplumber."""
        try:
            import pdfplumber
            text_content = ""
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Ekstrak teks
                    text = page.extract_text()
                    if text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{text}"
                    
                    # Ekstrak tabel jika ada
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            text_content += f"\n--- Table {table_num + 1} on Page {page_num + 1} ---\n"
                            for row in table:
                                text_content += "| " + " | ".join([str(cell) if cell else "" for cell in row]) + " |\n"
            
            return text_content.strip()
            
        except Exception as e:
            print(f"⚠️ PDF extraction failed: {e}")
            # Fallback ke PyPDF2 jika pdfplumber gagal
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(pdf_path)
                text_content = ""
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{text}"
                return text_content.strip()
            except Exception as e2:
                print(f"⚠️ PyPDF2 fallback also failed: {e2}")
                return None
    
    def _extract_docx_content(self, docx_path):
        """Ekstrak konten dari file DOCX."""
        try:
            from docx import Document
            doc = Document(docx_path)
            text_content = ""
            
            # Ekstrak paragraf
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Ekstrak tabel
            for table in doc.tables:
                text_content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        text_content += row_text + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            print(f"⚠️ DOCX extraction failed: {e}")
            return None
    
    def _extract_xlsx_content(self, xlsx_path):
        """Ekstrak konten dari file XLSX."""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(xlsx_path, read_only=True)
            text_content = ""
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                
                # Baca maksimal 100 baris pertama
                for row_idx, row in enumerate(ws.iter_rows(max_row=100)):
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                        else:
                            row_values.append("")
                    
                    # Hanya tambahkan baris yang tidak kosong
                    if any(val.strip() for val in row_values):
                        text_content += ", ".join(row_values) + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            print(f"⚠️ XLSX extraction failed: {e}")
            return None